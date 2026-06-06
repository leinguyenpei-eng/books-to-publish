"""
step3_finalize.py — Assemble + Cover + Export DOCX + Publish
═════════════════════════════════════════════════════════════
Source: tinh hoa từ book-automation (cover Pollinations+Pillow, publish Gumroad/Payhip)
         + book-final (DOCX KDP export) + telegram summary with cover image
- Viết Intro + Conclusion (Claude)
- Ghép tất cả chương thành DOCX chuẩn KDP 6×9"
- Tạo bìa: Pollinations AI background + Pillow text overlay
- Publish Gumroad + Payhip tự động
- Telegram thông báo cuối kèm ảnh bìa + links
- Email DOCX đầy đủ
"""
import json, os, re, sys, time, smtplib, requests, urllib.parse
from pathlib import Path
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

# ── Env ───────────────────────────────────────────────────────────────────────
CLAUDE_KEY     = os.environ.get("ANTHROPIC_API_KEY", "")
GMAIL_USER     = os.environ.get("GMAIL_USER", "")
GMAIL_APP_PASS = os.environ.get("GMAIL_APP_PASSWORD", "")
GMAIL_TO       = os.environ.get("GMAIL_TO", "") or GMAIL_USER
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT  = os.environ.get("TELEGRAM_CHAT_ID", "")
GUMROAD_KEY    = os.environ.get("GUMROAD_API_KEY", "")
PAYHIP_KEY     = os.environ.get("PAYHIP_API_KEY", "")

CLAUDE_MODEL = "claude-sonnet-4-6"
OUT_DIR      = Path("output")
OUT_DIR.mkdir(exist_ok=True)

# ── Load outline ──────────────────────────────────────────────────────────────
outline  = json.loads((OUT_DIR / "outline.json").read_text(encoding="utf-8"))
book_title = outline["final_title"]
subtitle   = outline.get("subtitle", "")
author      = outline.get("author", os.environ.get("AUTHOR_NAME", "Anonymous"))
author_style = os.environ.get("AUTHOR_STYLE", "Codie Sanchez (sharp, actionable) crossed with Malcolm Gladwell (story-driven, counterintuitive)")
niche      = outline.get("category", "self-help")
chapters   = outline["chapters"]

print(f"\n{'='*60}\n📚 Finalizing: {book_title}\n{'='*60}\n")

# ── Claude ────────────────────────────────────────────────────────────────────
WRITING_SYSTEM = """You are ghostwriting a premium Non-fiction book for Amazon KDP.
Your goal: write prose that feels 100% human — a real expert who has lived through this material.

═══ VOICE & RHYTHM ═══
- Vary sentence length deliberately: short punchy sentences. Then a longer one that builds a thought, adds texture, and earns its length before landing. Then short again.
- Add 1-2 natural imperfections per section: a rhetorical question the reader is already thinking, a brief digression that circles back, a moment of honest uncertainty.
- Write like Codie Sanchez at her most direct, or Malcolm Gladwell at his most curious — sharp, grounded, occasionally surprising.
- Use specific anecdotes, named places, real numbers. Not "a study found" but "a 2019 Stanford study of 400 middle managers found."
- First-person voice is allowed sparingly: "Here's what that actually means in practice."

═══ BANNED WORDS (AI fingerprints — never use these) ═══
Delve, Crucial, Leverage, Utilize, Synergy, Tapestry, Embark, Navigate/Navigating,
"It is important to note", "It's worth noting", "In today's fast-paced world",
"In conclusion", "Furthermore", "Moreover", "Overall", "In summary",
"This chapter will explore", "Let's dive into", "At the end of the day",
"Game-changer", "Paradigm shift", "Move the needle", "Cutting-edge", "Holistic"

═══ STRUCTURE ═══
- Fluent, premium American English only. Short paragraphs (3-5 sentences max).
- NO bullet lists — weave everything into flowing prose.
- ZERO repetition. Every argument needs a concrete example or specific data point.
- Start directly — no preamble, no "In this section we will..."
- NO summary paragraph at the end — end mid-momentum."""

def call_claude(prompt: str, max_tokens: int = 2000, retries: int = 3) -> str:
    for attempt in range(retries):
        try:
            r = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": CLAUDE_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json"
                },
                json={
                    "model": CLAUDE_MODEL, "max_tokens": max_tokens,
                    "system": WRITING_SYSTEM,
                    "messages": [{"role": "user", "content": prompt}]
                },
                timeout=180
            )
            r.raise_for_status()
            return r.json()["content"][0]["text"].strip()
        except Exception as e:
            wait = 30 * (attempt + 1)
            print(f"  ⚠ Claude attempt {attempt+1}: {e} — retry {wait}s")
            time.sleep(wait)
    raise RuntimeError("Claude failed")

# ── Write Intro ───────────────────────────────────────────────────────────────
print("✍ Writing Introduction...")
intro = call_claude(f"""Write the Introduction for "{book_title}: {subtitle}".

Chapters: {', '.join(c['title'] for c in chapters)}
Category: {niche}
Write in the style of: {author_style}

Requirements:
- 800-1,000 words
- Open with a specific story, scene, or provocative statement — not a generic claim
- Use the banned words list from your system prompt — no AI fingerprints
- Explain what readers will gain — without just listing chapters
- Add one moment of honest candor about what this book will NOT do
- End with momentum into Chapter 1
- NO bullet lists, flowing prose only""", max_tokens=1800)
print(f"  ✓ Intro: {len(intro.split())} words")

# ── Write Conclusion ──────────────────────────────────────────────────────────
print("✍ Writing Conclusion...")
conclusion = call_claude(f"""Write the Conclusion for "{book_title}: {subtitle}".

Chapters: {', '.join(c['title'] for c in chapters)}
Write in the style of: {author_style}

Requirements:
- 600-800 words
- DO NOT summarize chapters — that's boring and lazy
- Paint a vivid, specific picture of the reader's life AFTER applying everything
- Reference 2-3 specific insights naturally — woven in, not listed
- Add a rhetorical question that makes the reader pause
- End with a memorable, quotable final sentence — not a call to action cliché
- Inspire reflection AND action""", max_tokens=1500)
print(f"  ✓ Conclusion: {len(conclusion.split())} words")

# ── Load all chapters ─────────────────────────────────────────────────────────
print("📖 Loading all chapters...")
manuscript_chapters = []
total_words = len(intro.split()) + len(conclusion.split())

for ch in chapters:
    num  = ch["number"]
    path = OUT_DIR / f"chapter_{num:02d}.txt"
    if path.exists():
        content = path.read_text(encoding="utf-8")
        # Strip header if present
        if content.startswith("CHAPTER"):
            content = "\n".join(content.split("\n")[2:]).strip()
        wc = len(content.split())
        total_words += wc
        manuscript_chapters.append({"number": num, "title": ch["title"], "content": content})
        print(f"  ✓ Ch.{num}: {ch['title']} ({wc:,} words)")
    else:
        print(f"  ⚠ Ch.{num} missing — {path}")

print(f"\n📊 Total: {total_words:,} words across {len(manuscript_chapters)} chapters")

# ── Export DOCX (KDP 6×9") ───────────────────────────────────────────────────
print("\n📄 Exporting DOCX (KDP 6×9\")...")
doc = Document()

# Page setup
section = doc.sections[0]
section.page_width   = Inches(6)
section.page_height  = Inches(9)
section.left_margin  = section.right_margin  = Inches(0.875)
section.top_margin   = section.bottom_margin = Inches(0.75)

style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(22 if level == 1 else 15)

def add_body(text):
    for line in text.split("\n"):
        line = line.strip()
        if not line: continue
        if line.startswith("## "):
            add_heading(line[3:], level=2)
        elif line.startswith("# "):
            add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(11)

# Title page
add_heading(book_title, 1)
tp = doc.add_paragraph(subtitle); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap = doc.add_paragraph(f"\n{author}"); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Intro
add_heading("Introduction", 1)
add_body(intro)
doc.add_page_break()

# Chapters
for ch in manuscript_chapters:
    add_heading(f"Chapter {ch['number']}: {ch['title']}", 1)
    add_body(ch["content"])
    doc.add_page_break()

# Conclusion
add_heading("Conclusion", 1)
add_body(conclusion)

slug     = re.sub(r"[^\w]", "_", book_title)[:40].strip("_")
docx_path = OUT_DIR / f"{slug}_KDP.docx"
doc.save(docx_path)
print(f"  ✓ DOCX: {docx_path.name}")

# ── Generate Cover (Pollinations AI + Pillow) ─────────────────────────────────
print("\n🎨 Generating cover...")

COVER_W, COVER_H = 1600, 2560
THEMES = {
    "green":    {"bar": (29,158,117),  "title": (8,80,65),    "sub": (15,110,86),  "brand": (100,100,95), "overlay": (8,80,65)},
    "lavender": {"bar": (127,119,221), "title": (38,33,92),   "sub": (83,74,183),  "brand": (100,100,95), "overlay": (38,33,92)},
    "peach":    {"bar": (216,90,48),   "title": (113,43,19),  "sub": (153,60,29),  "brand": (100,100,95), "overlay": (113,43,19)},
    "blue":     {"bar": (24,95,165),   "title": (4,44,83),    "sub": (12,68,124),  "brand": (100,100,95), "overlay": (4,44,83)},
    "dark":     {"bar": (186,117,23),  "title": (255,255,255),"sub": (250,199,117),"brand": (180,180,175),"overlay": (20,20,18)},
    "pink":     {"bar": (212,83,126),  "title": (114,36,62),  "sub": (153,53,86),  "brand": (100,100,95), "overlay": (114,36,62)},
}

FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"
FONT_REG  = "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
FONT_BSANS= "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"

def load_font(path, size):
    try: return ImageFont.truetype(path, size)
    except: return ImageFont.load_default()

def wrap_text(draw, text, font, max_w):
    words, lines, cur = text.split(), [], ""
    for w in words:
        test = (cur + " " + w).strip()
        if draw.textbbox((0,0), test, font=font)[2] <= max_w: cur = test
        else:
            if cur: lines.append(cur)
            cur = w
    if cur: lines.append(cur)
    return lines

def draw_centered(draw, lines, font, color, y, lh, left, uw):
    for line in lines:
        tw = draw.textbbox((0,0), line, font=font)[2]
        draw.text((left + (uw-tw)//2, y), line, font=font, fill=color)
        y += lh
    return y

theme_name = outline.get("cover_theme", "blue")
art_prompt = outline.get("cover_art_prompt", "professional abstract background, calm colors, no text")
t = THEMES.get(theme_name, THEMES["blue"])
is_dark = theme_name == "dark"

# Get background from Pollinations (free, no API key)
full_prompt = (
    f"{art_prompt}, professional book cover background, "
    "NO text NO words NO letters NO numbers, high quality digital art, 300dpi"
)
encoded = urllib.parse.quote(full_prompt)
seed    = abs(hash(book_title)) % 99999
url     = f"https://image.pollinations.ai/prompt/{encoded}?width={COVER_W}&height={COVER_H}&nologo=true&enhance=true&seed={seed}"

bg_img = None
for attempt in range(3):
    try:
        r = requests.get(url, timeout=90, stream=True)
        if r.status_code == 200:
            bg_img = Image.open(BytesIO(r.content)).convert("RGBA")
            bg_img = bg_img.resize((COVER_W, COVER_H), Image.LANCZOS)
            print("  ✓ Background from Pollinations AI")
            break
    except Exception as e:
        print(f"  ⚠ Pollinations attempt {attempt+1}: {e}")
        time.sleep(5)

if bg_img is None:
    fallback = {"green":(232,244,240),"lavender":(238,237,254),"peach":(250,236,231),
                "blue":(230,241,251),"dark":(44,44,42),"pink":(251,234,240)}
    bg_img = Image.new("RGBA", (COVER_W, COVER_H), fallback.get(theme_name, (230,241,251)))
    print("  ⚠ Using solid color fallback")

# Gradient overlay
overlay = Image.new("RGBA", (COVER_W, COVER_H), (0,0,0,0))
ov_draw = ImageDraw.Draw(overlay)
ov_rgb  = t["overlay"]
for y_pos in range(COVER_H//2, COVER_H):
    alpha = int(200 * (y_pos - COVER_H//2) / (COVER_H//2))
    ov_draw.line([(0,y_pos),(COVER_W,y_pos)], fill=(*ov_rgb, min(alpha, 210)))
for y_pos in range(0, COVER_H//4):
    alpha = int(120 * (1 - y_pos/(COVER_H//4)))
    ov_draw.line([(0,y_pos),(COVER_W,y_pos)], fill=(*ov_rgb, alpha))

img  = Image.alpha_composite(bg_img, overlay).convert("RGB")
draw = ImageDraw.Draw(img)

# Accent bar
draw.rectangle([(0,0),(16,COVER_H)], fill=t["bar"])
left   = 96
usable = COVER_W - left - 80
title_color = (255,255,255) if is_dark else t["title"]
sub_color   = (220,220,220) if is_dark else t["sub"]
brand_color = (180,180,175) if is_dark else t["brand"]

# Title
f_title = load_font(FONT_BOLD, 150)
t_lines = wrap_text(draw, book_title.upper(), f_title, usable)
t_h     = len(t_lines) * 170
t_y     = COVER_H - t_h - 380
y = draw_centered(draw, t_lines, f_title, title_color, t_y, 170, left, usable)

# Divider
y += 40
f_div = load_font(FONT_REG, 52)
dw = draw.textbbox((0,0), "· · ·", font=f_div)[2]
draw.text((left + (usable-dw)//2, y), "· · ·", font=f_div, fill=t["bar"])
y += 80

# Subtitle
if subtitle:
    f_sub  = load_font(FONT_REG, 62)
    s_lines = wrap_text(draw, subtitle, f_sub, usable)
    y = draw_centered(draw, s_lines, f_sub, sub_color, y, 78, left, usable)

# Author
f_auth = load_font(FONT_BSANS, 46)
bw = draw.textbbox((0,0), author.upper(), font=f_auth)[2]
draw.line([(left+60, COVER_H-145),(COVER_W-60, COVER_H-145)], fill=(*t["bar"],80), width=2)
draw.text((left + (usable-bw)//2, COVER_H-120), author.upper(), font=f_auth, fill=brand_color)

cover_path = OUT_DIR / f"{slug}_cover.png"
img.save(str(cover_path), "PNG", dpi=(300,300))
print(f"  ✓ Cover: {cover_path.name} ({cover_path.stat().st_size//1024}KB)")

# ── Gmail helpers ─────────────────────────────────────────────────────────────
def send_email(subject, body, attachment=None):
    if not GMAIL_USER or not GMAIL_APP_PASS:
        print("  ⚠ Gmail not configured"); return
    try:
        msg = MIMEMultipart()
        msg["From"] = GMAIL_USER
        msg["To"]   = GMAIL_TO or GMAIL_USER
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain", "utf-8"))
        if attachment and Path(attachment).exists():
            with open(attachment, "rb") as f:
                part = MIMEBase("application","octet-stream")
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={Path(attachment).name}")
            msg.attach(part)
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as s:
            s.login(GMAIL_USER, GMAIL_APP_PASS)
            s.send_message(msg)
        print(f"  ✉ Email sent: {subject[:60]}")
    except Exception as e:
        print(f"  ⚠ Email: {e}")

# ── Telegram ──────────────────────────────────────────────────────────────────
def telegram_send(msg, image=None):
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT: return
    try:
        if image and Path(image).exists():
            with open(image, "rb") as f:
                requests.post(
                    f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendPhoto",
                    data={"chat_id": TELEGRAM_CHAT, "caption": msg[:1024], "parse_mode": "HTML"},
                    files={"photo": ("cover.png", f, "image/png")},
                    timeout=30
                )
        else:
            requests.post(
                f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage",
                json={"chat_id": TELEGRAM_CHAT, "text": msg, "parse_mode": "HTML"},
                timeout=15
            )
    except Exception as e:
        print(f"  ⚠ Telegram: {e}")

# ── Publish Gumroad ───────────────────────────────────────────────────────────
gumroad_url = ""
if GUMROAD_KEY:
    print("\n🚀 Publishing to Gumroad...")
    try:
        pricing = outline.get("pricing", {})
        price   = int(pricing.get("gumroad", {}).get("price", 9.99) * 100)
        r = requests.post(
            "https://api.gumroad.com/v2/products",
            data={
                "access_token": GUMROAD_KEY,
                "name": f"{book_title}: {subtitle}",
                "description": f"{outline.get('description','')}\n\n✨ Instant digital download",
                "price": price, "currency": "usd", "published": "true"
            },
            timeout=30
        )
        pid = r.json().get("product", {}).get("id")
        gumroad_url = r.json().get("product", {}).get("short_url", "")
        if pid:
            with open(docx_path, "rb") as f:
                requests.post(
                    f"https://api.gumroad.com/v2/products/{pid}/files",
                    data={"access_token": GUMROAD_KEY},
                    files={"file": (docx_path.name, f)},
                    timeout=60
                )
            if cover_path.exists():
                with open(cover_path, "rb") as f:
                    requests.post(
                        f"https://api.gumroad.com/v2/products/{pid}",
                        data={"access_token": GUMROAD_KEY},
                        files={"preview_url": (cover_path.name, f, "image/png")},
                        timeout=30
                    )
            print(f"  ✓ Gumroad: {gumroad_url}")
    except Exception as e:
        print(f"  ⚠ Gumroad: {e}")
else:
    print("  ⚠ GUMROAD_API_KEY not set — skipping")

# ── Publish Payhip ────────────────────────────────────────────────────────────
payhip_url = ""
if PAYHIP_KEY:
    print("🚀 Publishing to Payhip...")
    try:
        pricing = outline.get("pricing", {})
        price   = pricing.get("payhip", {}).get("price", 9.99)
        r = requests.post(
            "https://payhip.com/api/v1/product",
            headers={"payhip-api-key": PAYHIP_KEY},
            json={
                "title": f"{book_title}: {subtitle}",
                "description": outline.get("description", ""),
                "price": price, "type": "digital_download"
            },
            timeout=30
        )
        pid = r.json().get("data", {}).get("link")
        if pid:
            with open(docx_path, "rb") as f:
                requests.post(
                    f"https://payhip.com/api/v1/product/{pid}/file",
                    headers={"payhip-api-key": PAYHIP_KEY},
                    files={"file": (docx_path.name, f)},
                    timeout=60
                )
            payhip_url = f"https://payhip.com/b/{pid}"
            print(f"  ✓ Payhip: {payhip_url}")
    except Exception as e:
        print(f"  ⚠ Payhip: {e}")
else:
    print("  ⚠ PAYHIP_API_KEY not set — skipping")

# ── Email final DOCX ──────────────────────────────────────────────────────────
links_text = ""
if gumroad_url: links_text += f"\nGumroad: {gumroad_url}"
if payhip_url:  links_text += f"\nPayhip:  {payhip_url}"

send_email(
    subject=f"🎉 COMPLETE: {book_title} ({total_words:,} words)",
    body=(
        f"Your book is complete!\n\n"
        f"Title: {book_title}\n"
        f"Subtitle: {subtitle}\n"
        f"Author: {author}\n"
        f"Total words: {total_words:,}\n"
        f"Chapters: {len(manuscript_chapters)}\n"
        f"{links_text}\n\n"
        f"KDP: Upload DOCX manually at kdp.amazon.com\n"
        f"Etsy: Upload PDF manually with cover image\n\n"
        f"DOCX attached — ready for KDP upload."
    ),
    attachment=docx_path
)

# ── Etsy upload guide (text file) ─────────────────────────────────────────────
etsy_tags = outline.get("etsy_tags", [])
kdp_kw    = outline.get("keywords_kdp", [])
pricing   = outline.get("pricing", {})

guide = f"""📚 UPLOAD GUIDE — {book_title}
{'='*60}

GUMROAD: {gumroad_url or 'Log in → New Product → upload DOCX'}
PAYHIP:  {payhip_url  or 'Log in → Add Product → upload DOCX'}

KDP (Amazon): kdp.amazon.com
  → Bookshelf → Create → eBook
  Title: {book_title}
  Subtitle: {subtitle}
  Author: {author}
  Description: {outline.get('description','')}
  Keywords: {', '.join(kdp_kw)}
  Price: ${pricing.get('kdp',{}).get('price',9.99)}
  Upload file: {docx_path.name}

ETSY: etsy.com/sell → Add listing
  Title: {book_title}: {subtitle}
  Tags: {', '.join(etsy_tags)}
  Price: ${pricing.get('etsy',{}).get('price',6.99)}
  Upload: PDF version + cover image

Total words: {total_words:,}
Chapters: {len(manuscript_chapters)}
"""

(OUT_DIR / "UPLOAD_GUIDE.txt").write_text(guide, encoding="utf-8")
print("\n✓ Upload guide saved to output/UPLOAD_GUIDE.txt")

# ── Final Telegram ────────────────────────────────────────────────────────────
links_html = ""
if gumroad_url: links_html += f"\n🛒 Gumroad: {gumroad_url}"
if payhip_url:  links_html += f"\n🛒 Payhip: {payhip_url}"

telegram_send(
    f"🎉 <b>BOOK COMPLETE!</b>\n\n"
    f"<b>{book_title}</b>\n"
    f"<i>{subtitle}</i>\n\n"
    f"📊 {total_words:,} words · {len(manuscript_chapters)} chapters\n"
    f"✍ Author: {author}\n"
    f"📧 Full DOCX sent to Gmail{links_html}\n\n"
    f"📤 KDP + Etsy: see UPLOAD_GUIDE.txt in GitHub",
    image=cover_path
)

print(f"\n{'='*60}")
print(f"🎉 COMPLETE: {book_title}")
print(f"   Words:  {total_words:,}")
print(f"   DOCX:   {docx_path.name}")
print(f"   Cover:  {cover_path.name}")
if gumroad_url: print(f"   Gumroad: {gumroad_url}")
if payhip_url:  print(f"   Payhip:  {payhip_url}")
print(f"{'='*60}\n")
