"""
YouTube → Book Pipeline (FINAL — Combined Best Version)
═══════════════════════════════════════════════════════
Step 1  Fetch YouTube transcripts (6 fallback methods)
Step 2  Gemini synthesizes outline (12 chapters, large context)
Step 3  Claude Sonnet writes chapters in 3 sections each
Step 4  Export DOCX — KDP 6×9" format, EN + VI
Step 5  Generate book cover (Imagen 3 + Pillow overlay)
Step 6  Publish → Gumroad + Payhip + Beacons + Fourthwall
Step 7  Facebook posts (5 types × EN + VI = 10 posts)
Step 8  TikTok / Instagram / LinkedIn / Twitter content
Step 9  Email campaign (Mailchimp + SendGrid + HTML draft)
Step 10 Trigger n8n webhook (optional Canva automation)
"""

import os, re, time, json, yaml, glob, sys
import requests
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Keys (all from GitHub Secrets — never hardcoded) ──────────────────────────
GEMINI_KEY     = os.environ.get("GEMINI_API_KEY", "")
CLAUDE_KEY     = os.environ.get("ANTHROPIC_API_KEY", "")
GUMROAD_KEY    = os.environ.get("GUMROAD_API_KEY", "")
PAYHIP_KEY     = os.environ.get("PAYHIP_API_KEY", "")
BEACONS_KEY    = os.environ.get("BEACONS_API_KEY", "")
FOURTHWALL_KEY = os.environ.get("FOURTHWALL_API_KEY", "")

GEMINI_URL   = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
CLAUDE_URL   = "https://api.anthropic.com/v1/messages"
CLAUDE_MODEL = "claude-sonnet-4-6"

OUTPUT_DIR = Path(__file__).parent.parent / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ── Logging ───────────────────────────────────────────────────────────────────
def log(msg):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}", flush=True)


# ── AI callers ────────────────────────────────────────────────────────────────
GEMINI_MODELS = [
    "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent",
    "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-8b:generateContent",
]

def call_gemini(prompt: str, retries=4, max_tokens=2048) -> str:
    """Smart retry: 429 → longer wait + fallback model."""
    for attempt in range(retries):
        model = GEMINI_MODELS[min(attempt, len(GEMINI_MODELS)-1)]
        try:
            r = requests.post(
                f"{model}?key={GEMINI_KEY}",
                json={"contents": [{"parts": [{"text": prompt}]}],
                      "generationConfig": {"temperature": 0.7, "maxOutputTokens": max_tokens}},
                timeout=90)
            if r.status_code == 429:
                wait = [60, 90, 120, 180][min(attempt, 3)]
                log(f"  ⚠ Gemini 429 rate limit (attempt {attempt+1}) — waiting {wait}s...")
                time.sleep(wait)
                continue
            r.raise_for_status()
            return r.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
        except Exception as e:
            wait = 30 * (attempt + 1)
            log(f"  ⚠ Gemini attempt {attempt+1}: {e} — retry {wait}s")
            time.sleep(wait)
    raise RuntimeError("Gemini failed after retries — check GEMINI_API_KEY")


def call_claude(prompt: str, system: str = "", retries=3, max_tokens=2048) -> str:
    if not CLAUDE_KEY:
        raise RuntimeError("No ANTHROPIC_API_KEY set")
    for attempt in range(retries):
        try:
            body = {
                "model": CLAUDE_MODEL,
                "max_tokens": max_tokens,
                "temperature": 0.85,
                "messages": [{"role": "user", "content": prompt}],
            }
            if system:
                body["system"] = system
            r = requests.post(
                CLAUDE_URL,
                headers={
                    "x-api-key": CLAUDE_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json=body,
                timeout=90)
            r.raise_for_status()
            return r.json()["content"][0]["text"].strip()
        except Exception as e:
            wait = 20 * (attempt + 1)
            log(f"  ⚠ Claude attempt {attempt+1}: {e} — retry {wait}s")
            time.sleep(wait)
    raise RuntimeError("Claude failed after retries")


def call_ai(prompt: str, prefer: str = "claude", system: str = "") -> str:
    """Smart router: prefer=claude → Claude first, fallback Gemini (and vice versa)"""
    if prefer == "claude":
        try:
            return call_claude(prompt, system=system)
        except Exception as e:
            log(f"  ↩ Claude failed ({e}), falling back to Gemini...")
            return call_gemini(prompt)
    else:
        try:
            return call_gemini(prompt)
        except Exception as e:
            log(f"  ↩ Gemini failed ({e}), falling back to Claude...")
            return call_claude(prompt, system=system)


# ── Step 1: Transcripts ───────────────────────────────────────────────────────

def _ai_research_fallback(config: dict) -> str:
    """Khi IP bị block, dùng Claude/Gemini tự research topic."""
    title = config.get("title", "Untitled")
    niche = config.get("niche", "Self-help")
    desc  = config.get("description", "")
    urls  = config.get("youtube_urls", [])
    url_context = f"\nDrawing from {len(urls)} expert videos on this topic." if urls else ""
    prompt = f"""You are a research assistant for a book.

Book: "{title}" — {niche}
Description: {desc}{url_context}

Write 2500 words of comprehensive research content covering:
1. Key theories and frameworks from leading experts
2. Scientific findings and statistics
3. Real-world case studies
4. Common challenges and proven solutions  
5. Actionable strategies with evidence
6. Expert insights and counterintuitive findings

Be specific with names, studies, and examples. Write as flowing research notes."""

    try:
        return call_claude(prompt, max_tokens=4000)
    except Exception:
        try:
            return call_gemini(prompt, max_tokens=4000)
        except Exception:
            return f'Research for "{title}" ({niche}): {desc}'


def gather_sources(urls: list) -> str:
    """
    Fetch transcripts với fast-fail:
    Nếu 3 URL đầu đều fail → IP bị block → skip hết, dùng AI research fallback.
    """
    sys.path.insert(0, str(Path(__file__).parent))
    from transcript_fetcher import fetch_transcript

    log(f"📥 Fetching {len(urls)} transcripts...")
    parts = []
    consecutive_fails = 0

    for i, url in enumerate(urls, 1):
        log(f"  [{i}/{len(urls)}] {url[:65]}")
        t = fetch_transcript(url)
        if t and len(t.split()) > 50:
            parts.append(f"=== SOURCE {i} ===\n{t}")
            consecutive_fails = 0
        else:
            consecutive_fails += 1
            if consecutive_fails >= 3 and i <= 5:
                log(f"  → {consecutive_fails} consecutive fails — IP likely blocked, skipping remaining URLs")
                log(f"  → Will use Claude AI research fallback instead")
                break
        time.sleep(1)

    combined = "\n\n".join(parts)
    if combined:
        log(f"✓ Total: {len(combined.split())} words from {len(parts)}/{len(urls)} sources")
    else:
        log("⚠ No transcripts fetched — switching to Claude AI research mode")
        combined = _ai_research_fallback(config_fallback)
        log(f"✓ AI research: {len(combined.split())} words generated")
    return combined


# ── Step 2: Outline ───────────────────────────────────────────────────────────
def build_outline(raw: str, config: dict) -> dict:
    log("🧠 Building outline (Gemini — large context)...")
    title = config.get("title", "Untitled")
    niche = config.get("niche", "Self-help")

    prompt = f"""Analyze these YouTube transcripts and create a detailed book outline for "{title}" ({niche}).

TRANSCRIPTS:
{raw[:80000]}

Return ONLY valid JSON, no markdown fences, no explanation:
{{
  "core_thesis": "2-3 sentences summarizing the book's main argument",
  "target_reader": "1 sentence describing the ideal reader",
  "introduction_hook": "3-4 sentence compelling opening hook",
  "conclusion_message": "2-3 sentence powerful final message",
  "chapters": [
    {{
      "number": 1,
      "title": "Compelling chapter title",
      "premise": "2 sentences explaining what this chapter argues",
      "key_points": ["point 1", "point 2", "point 3", "point 4"],
      "stories_or_examples": ["specific example from transcripts"],
      "data_or_quotes": ["specific stat or quote from sources"]
    }}
  ]
}}
Generate exactly 12 chapters. ONLY JSON output."""

    raw_out = call_gemini(prompt, max_tokens=4096)
    raw_out = re.sub(r"^```json\s*", "", raw_out, flags=re.MULTILINE)
    raw_out = re.sub(r"```\s*$", "", raw_out, flags=re.MULTILINE).strip()

    try:
        outline = json.loads(raw_out)
        log(f"✓ Outline: {len(outline['chapters'])} chapters")
        return outline
    except Exception as e:
        log(f"⚠ JSON parse failed ({e}) — using fallback outline")
        return {
            "core_thesis": title, "target_reader": "General readers",
            "introduction_hook": "Begin.", "conclusion_message": "End.",
            "chapters": [
                {"number": i, "title": f"Chapter {i}", "premise": "Key ideas.",
                 "key_points": ["P1", "P2", "P3"], "stories_or_examples": [], "data_or_quotes": []}
                for i in range(1, 13)
            ]
        }


# ── Step 3: Write chapters ────────────────────────────────────────────────────
def write_section(ch: dict, config: dict, lang: str, s_num: int, prev: str) -> str:
    li    = "Write in English." if lang == "en" else "Viết hoàn toàn bằng Tiếng Việt tự nhiên, không dịch cứng."
    title = config.get("title", "Untitled")
    niche = config.get("niche", "Self-help")
    kp    = "\n".join(f"- {p}" for p in ch.get("key_points", []))
    st    = "\n".join(f"- {s}" for s in ch.get("stories_or_examples", []))
    da    = "\n".join(f"- {d}" for d in ch.get("data_or_quotes", []))

    tasks = {
        1: "Write the OPENING: strong hook + introduce the chapter premise. End mid-thought naturally to flow into next section.",
        2: f"CONTINUE seamlessly from: '...{prev[-200:]}'\nWrite the MIDDLE: deep dive into the key points, weave in stories and data naturally.",
        3: f"CONTINUE seamlessly from: '...{prev[-200:]}'\nWrite the CLOSING: synthesize key ideas, give 1 concrete action step, bridge to next chapter.",
    }

    prompt = f"""You are ghostwriting a {niche} book titled "{title}".
{li}

CHAPTER {ch['number']}: {ch['title']}
Premise: {ch.get('premise', '')}
Key points:
{kp}
Stories/Examples to weave in:
{st}
Data/Quotes to reference:
{da}

TASK: {tasks[s_num]}

Writing rules (KDP nonfiction standard):
- Conversational but authoritative voice
- Short paragraphs (3-5 sentences max)
- Use ## Subheading where natural to break sections
- No bullet lists — weave everything into flowing prose
- ~600-700 words for this section
- Start directly with content — no labels, no "Section X:" headers

Write now:"""

    return call_ai(prompt, prefer="claude")


def write_chapter(ch: dict, config: dict, lang: str) -> str:
    log(f"  ✍ Ch.{ch['number']}: {ch['title']} ({lang.upper()})")
    full = ""
    for s_num in range(1, 4):
        sec = write_section(ch, config, lang, s_num, full)
        full += "\n\n" + sec
        log(f"    §{s_num}/3 — {len(sec.split())} words")
        time.sleep(3)
    log(f"  ✓ Ch.{ch['number']} done: {len(full.split())} words")
    return full.strip()


def write_intro_conclusion(outline: dict, config: dict, lang: str) -> tuple:
    li    = "Write in English." if lang == "en" else "Viết hoàn toàn bằng Tiếng Việt tự nhiên."
    title = config.get("title", "Untitled")
    niche = config.get("niche", "Self-help")
    ch_list = "\n".join(f"- Chapter {c['number']}: {c['title']}" for c in outline.get("chapters", []))

    log(f"  ✍ Introduction ({lang.upper()})...")
    intro = call_ai(f"""Write a compelling Introduction for "{title}" ({niche} book).
{li}
Hook: {outline.get('introduction_hook', '')}
Core thesis: {outline.get('core_thesis', '')}
Target reader: {outline.get('target_reader', '')}
Chapters overview:
{ch_list}

800-1000 words. KDP nonfiction style. Preview what readers will gain from each chapter.
Start directly — no "Introduction" header needed:""", prefer="claude")

    time.sleep(3)

    log(f"  ✍ Conclusion ({lang.upper()})...")
    conclusion = call_ai(f"""Write a powerful Conclusion for "{title}" ({niche} book).
{li}
Final message: {outline.get('conclusion_message', '')}
Core thesis: {outline.get('core_thesis', '')}

600-800 words. Inspiring + actionable. Summarize transformation, give final call to action.
Start directly:""", prefer="claude")

    return intro, conclusion


# ── Step 4: DOCX (KDP 6×9") ──────────────────────────────────────────────────
def build_docx(ms: dict, config: dict, lang: str) -> Path:
    title  = config.get("title", "Untitled")
    author = config.get("author", "Unknown Author")
    suffix = "EN" if lang == "en" else "VI"
    path   = OUTPUT_DIR / f"{title} ({suffix}).docx"

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Inches(6);  sec.page_height  = Inches(9)
    sec.left_margin = sec.right_margin  = Inches(0.75)
    sec.top_margin  = sec.bottom_margin = Inches(0.75)

    sty = doc.styles["Normal"]
    sty.font.name = "Times New Roman"; sty.font.size = Pt(12)
    sty.paragraph_format.line_spacing = Pt(18)
    sty.paragraph_format.space_after  = Pt(0)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(20 if level == 1 else 14)
        p.paragraph_format.space_before = Pt(24 if level == 1 else 18)
        p.paragraph_format.space_after  = Pt(12)

    def add_body(text):
        # Split on ## subheadings
        parts = re.split(r"\n##\s+(.+?)(?:\n|$)", text)
        i = 0
        while i < len(parts):
            for para in parts[i].strip().split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(para)
                run.font.name = "Times New Roman"; run.font.size = Pt(12)
            i += 1
            if i < len(parts):
                add_heading(parts[i].strip(), level=2)
                i += 1

    # ── Title page ──
    doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title); tr.bold = True
    tr.font.size = Pt(24); tr.font.name = "Times New Roman"
    doc.add_paragraph()
    subtitle = config.get("subtitle", "")
    if subtitle:
        sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run(subtitle).font.size = Pt(14)
    ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.add_run(f"by {author}").font.size = Pt(14)
    doc.add_page_break()

    # ── Copyright page ──
    cp = doc.add_paragraph()
    cp.add_run(f"Copyright © {datetime.now().year} {author}. All rights reserved.\n"
               f"No part of this publication may be reproduced without written permission.")
    cp.runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── Content ──
    add_heading("Introduction"); add_body(ms["introduction"]); doc.add_page_break()
    for ch in ms["chapters"]:
        add_heading(f"Chapter {ch['number']}: {ch['title']}")
        add_body(ch["content"])
        doc.add_page_break()
    add_heading("Conclusion"); add_body(ms["conclusion"])

    doc.save(str(path))
    size_kb = path.stat().st_size // 1024
    log(f"✓ DOCX saved: {path.name} ({size_kb} KB)")
    return path


# ── Step 5: Cover ─────────────────────────────────────────────────────────────
def make_cover(config: dict):
    try:
        sys.path.insert(0, str(Path(__file__).parent))
        from cover_generator import make_cover as _make
        return _make(config)
    except Exception as e:
        log(f"⚠ Cover skipped: {e}")
        return None


# ── Step 6: Publish ───────────────────────────────────────────────────────────
def publish_gumroad(path: Path, config: dict, cover: Path = None) -> str:
    if not GUMROAD_KEY:
        log("⚠ No GUMROAD_API_KEY"); return ""
    title = config.get("title", "Untitled")
    price = int(float(config.get("price_usd", 9.99)) * 100)
    log(f"📤 Publishing to Gumroad: {title}")
    try:
        r = requests.post("https://api.gumroad.com/v2/products", data={
            "access_token": GUMROAD_KEY, "name": title, "price": price,
            "description": config.get("description", ""), "published": "true"}, timeout=30)
        data = r.json()
        if not data.get("success"):
            log(f"  ⚠ Gumroad error: {data}"); return ""
        pid = data["product"]["id"]
        with open(path, "rb") as f:
            requests.put(f"https://api.gumroad.com/v2/products/{pid}/files",
                         data={"access_token": GUMROAD_KEY},
                         files={"file": (path.name, f)}, timeout=60)
        if cover and cover.exists():
            with open(cover, "rb") as f:
                requests.put(f"https://api.gumroad.com/v2/products/{pid}/cover_image",
                             data={"access_token": GUMROAD_KEY},
                             files={"cover_image": (cover.name, f, "image/jpeg")}, timeout=30)
        url = data["product"].get("short_url", "")
        log(f"  ✓ Gumroad: {url}"); return url
    except Exception as e:
        log(f"  ⚠ Gumroad error: {e}"); return ""


def publish_payhip(path: Path, config: dict, cover: Path = None) -> str:
    if not PAYHIP_KEY:
        log("⚠ No PAYHIP_API_KEY"); return ""
    log(f"📤 Payhip: {config.get('title')} — add endpoint when Payhip releases public API")
    return ""


def publish_beacons(path: Path, config: dict) -> str:
    if not BEACONS_KEY:
        log("⚠ No BEACONS_API_KEY"); return ""
    log(f"📤 Publishing to Beacons: {config.get('title')}")
    try:
        r = requests.post("https://api.beacons.ai/v1/products",
            headers={"Authorization": f"Bearer {BEACONS_KEY}", "Content-Type": "application/json"},
            json={"title": config.get("title", ""), "description": config.get("description", ""),
                  "price": float(config.get("price_usd", 9.99)), "currency": "USD",
                  "type": "digital_download"}, timeout=30)
        data = r.json()
        if data.get("id"):
            url = data.get("url", "")
            log(f"  ✓ Beacons: {url}"); return url
        log(f"  ⚠ Beacons: {data}")
    except Exception as e:
        log(f"  ⚠ Beacons error: {e}")
    return ""


def publish_fourthwall(path: Path, config: dict) -> str:
    if not FOURTHWALL_KEY:
        log("⚠ No FOURTHWALL_API_KEY"); return ""
    log(f"📤 Publishing to Fourthwall: {config.get('title')}")
    try:
        r = requests.post("https://api.fourthwall.com/v1/products",
            headers={"Authorization": f"Bearer {FOURTHWALL_KEY}", "Content-Type": "application/json"},
            json={"name": config.get("title", ""), "description": config.get("description", ""),
                  "price_in_cents": int(float(config.get("price_usd", 9.99)) * 100),
                  "product_type": "digital"}, timeout=30)
        data = r.json()
        if data.get("id"):
            url = data.get("url", "")
            log(f"  ✓ Fourthwall: {url}"); return url
        log(f"  ⚠ Fourthwall: {data}")
    except Exception as e:
        log(f"  ⚠ Fourthwall error: {e}")
    return ""


# ── Step 7-9: Marketing ───────────────────────────────────────────────────────
def run_marketing(config: dict, buy_links: dict, cover: Path):
    # Facebook posts (5 types × EN + VI)
    try:
        sys.path.insert(0, str(Path(__file__).parent))
        from social_publisher import run_social_publisher
        run_social_publisher(config, buy_links, cover)
    except Exception as e:
        log(f"⚠ Facebook posts skipped: {e}")

    # TikTok / Instagram / LinkedIn / Twitter
    try:
        from multi_social import run_multi_social
        run_multi_social(config, buy_links, cover)
    except Exception as e:
        log(f"⚠ Multi-social skipped: {e}")

    # Email campaign
    try:
        from email_marketing import run_email_campaign
        run_email_campaign(config, buy_links, cover)
    except Exception as e:
        log(f"⚠ Email campaign skipped: {e}")


# ── Main orchestrator ─────────────────────────────────────────────────────────
config_fallback = {}  # used in gather_sources fallback

def build_book(config: dict):
    global config_fallback
    config_fallback = config

    title = config.get("title", "Untitled")
    urls  = config.get("youtube_urls", [])
    langs = config.get("languages", ["en", "vi"])

    log(f"\n{'='*60}\n📚 {title}\n   {len(urls)} sources · {', '.join(langs)}\n{'='*60}\n")

    # 1. Transcripts
    raw = gather_sources(urls)

    # 2. Outline
    outline = build_outline(raw, config)
    outline_path = OUTPUT_DIR / f"{title}_outline.json"
    outline_path.write_text(json.dumps(outline, indent=2, ensure_ascii=False), encoding="utf-8")
    log(f"✓ Outline saved: {outline_path.name}")

    # 3+4. Write manuscript + export DOCX for each language
    en_docx = None
    for lang in langs:
        flag = "🇺🇸 EN" if lang == "en" else "🇻🇳 VI"
        log(f"\n{flag} Writing manuscript...\n")

        intro, conclusion = write_intro_conclusion(outline, config, lang)
        ms = {"introduction": intro, "conclusion": conclusion, "chapters": []}

        for ch in outline["chapters"]:
            content = write_chapter(ch, config, lang)
            ms["chapters"].append({"number": ch["number"], "title": ch["title"], "content": content})
            time.sleep(2)

        total_words = len((intro + conclusion + " ".join(c["content"] for c in ms["chapters"])).split())
        log(f"\n✓ {flag}: {total_words:,} words total")

        docx_path = build_docx(ms, config, lang)
        if lang == "en":
            en_docx = docx_path

    # 5. Cover
    cover = make_cover(config)

    # 6. Publish
    buy_links = {}
    if en_docx:
        buy_links["Gumroad"]    = publish_gumroad(en_docx, config, cover)
        buy_links["Payhip"]     = publish_payhip(en_docx, config, cover)
        buy_links["Beacons"]    = publish_beacons(en_docx, config)
        buy_links["Fourthwall"] = publish_fourthwall(en_docx, config)

    # Etsy publish
    try:
        sys.path.insert(0, str(Path(__file__).parent))
        from publish_etsy import publish_to_etsy
        etsy_url = publish_to_etsy(en_docx, config, cover)
        if etsy_url:
            buy_links["Etsy"] = etsy_url
    except Exception as e:
        log(f"⚠ Etsy skipped: {e}")

    # KDP + upload guide for other platforms
    try:
        from kdp_guide import generate_upload_guide
        import json as _json
        pricing_path = Path(__file__).parent.parent / "config" / "pricing.json"
        pricing = {}
        if pricing_path.exists():
            niche_map = {"Self-help": "ebook", "Business": "ebook", "Finance": "spreadsheet",
                         "Health": "workbook", "Productivity": "planner"}
            btype = niche_map.get(config.get("niche", "Self-help"), "ebook")
            all_pricing = _json.loads(pricing_path.read_text())
            pricing = all_pricing.get("book_types", {}).get(btype, {}).get("pricing", {})
        generate_upload_guide(en_docx, config, cover, buy_links, pricing)
    except Exception as e:
        log(f"⚠ KDP guide skipped: {e}")

    # 7-9. Marketing
    run_marketing(config, buy_links, cover)

    # 10. n8n trigger
    if en_docx:
        try:
            from n8n_trigger import trigger_n8n
            trigger_n8n(config, en_docx, buy_links)
        except Exception as e:
            log(f"⚠ n8n trigger skipped: {e}")

    # Summary
    log(f"\n{'='*60}")
    log(f"🎉 DONE: {title}")
    log(f"   Output folder: {OUTPUT_DIR}/")
    active_links = {k: v for k, v in buy_links.items() if v}
    if active_links:
        for k, v in active_links.items():
            log(f"   🔗 {k}: {v}")
    else:
        log("   ⚠ No publish links (check API keys in GitHub Secrets)")
    log(f"{'='*60}\n")


def main():
    files = sorted(glob.glob("books/*.yml"))
    if not files:
        log("❌ No .yml files in books/"); return
    log(f"Found {len(files)} book(s)")
    for f in files:
        config = yaml.safe_load(open(f, encoding="utf-8"))
        build_book(config)


if __name__ == "__main__":
    main()
